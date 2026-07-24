# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Times New Roman"

RPR_ORDER = ['rStyle', 'rFonts', 'b', 'bCs', 'i', 'iCs', 'caps', 'smallCaps', 'strike',
             'dstrike', 'outline', 'shadow', 'emboss', 'imprint', 'noProof', 'snapToGrid',
             'vanish', 'webHidden', 'color', 'spacing', 'w', 'kern', 'position', 'sz', 'szCs',
             'highlight', 'u', 'effect', 'bdr', 'shd', 'fitText', 'vertAlign', 'rtl', 'cs',
             'em', 'lang', 'eastAsianLayout', 'specVanish', 'oMath']

TBLPR_ORDER = ['tblStyle', 'tblpPr', 'tblOverlap', 'bidiVisual', 'tblStyleRowBandSize',
               'tblStyleColBandSize', 'tblW', 'jc', 'tblCellSpacing', 'tblInd', 'tblBorders',
               'shd', 'tblLayout', 'tblCellMar', 'tblLook', 'tblCaption', 'tblDescription']


def insert_in_order(parent, new_el, order_list):
    tag = new_el.tag.split('}')[-1]
    try:
        new_idx = order_list.index(tag)
    except ValueError:
        parent.append(new_el)
        return
    for child in parent:
        child_tag = child.tag.split('}')[-1]
        if child_tag in order_list and order_list.index(child_tag) > new_idx:
            child.addprevious(new_el)
            return
    parent.append(new_el)


def set_font(run, size=13, bold=False, italic=False, color=None, underline=False):
    run.font.name = FONT
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        insert_in_order(rPr, rFonts, RPR_ORDER)
    rFonts.set(qn('w:ascii'), FONT)
    rFonts.set(qn('w:hAnsi'), FONT)
    rFonts.set(qn('w:eastAsia'), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_run(p, text, **kw):
    r = p.add_run(text)
    set_font(r, **kw)
    return r


def add_rich_paragraph(doc_or_cell, text, size=13, align=None, space_after=6, space_before=0,
                        first_line_indent=None, line_spacing=1.3, bold=False):
    p = doc_or_cell.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = line_spacing
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    add_run(p, text, size=size, bold=bold)
    return p


def set_cell_no_margins(cell, top=40, bottom=40, left=80, right=80):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement('w:tcMar')
    for tag, val in (('top', top), ('start', left), ('bottom', bottom), ('end', right)):
        el = OxmlElement(f'w:{tag}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        mar.append(el)
    tcPr.append(mar)


def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'nil')
        borders.append(el)
    insert_in_order(tblPr, borders, TBLPR_ORDER)


def set_table_borders(table, size=4, color='000000'):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(size))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        borders.append(el)
    insert_in_order(tblPr, borders, TBLPR_ORDER)


def set_col_widths(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for idx, w in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(w)
    for idx, w in enumerate(widths_cm):
        table.columns[idx].width = Cm(w)


def new_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    style = doc.styles['Normal']
    style.font.name = FONT
    style.font.size = Pt(13)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT)
    return doc


def add_header_block(doc):
    table = doc.add_table(rows=1, cols=2)
    remove_table_borders(table)
    set_col_widths(table, [6.0, 10.0])
    left, right = table.rows[0].cells
    set_cell_no_margins(left)
    set_cell_no_margins(right)

    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Công ty TNHH Giải pháp Chuyển đổi số NEXUS", size=12, bold=True)
    p2 = left.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(30)
    add_run(p2, "Số: ........................", size=13)

    q = right.paragraphs[0]
    q.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(q, "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", size=13, bold=True)
    q2 = right.add_paragraph()
    q2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(q2, "Độc lập - Tự do - Hạnh phúc", size=13, bold=True, underline=True)
    q3 = right.add_paragraph()
    q3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    q3.paragraph_format.space_before = Pt(30)
    add_run(q3, "Thành phố Hà Nội, ngày 23 tháng 07 năm 2026", size=13, italic=True)
    return table


doc = new_doc()
add_header_block(doc)

title1 = doc.add_paragraph()
title1.alignment = WD_ALIGN_PARAGRAPH.CENTER
title1.paragraph_format.space_before = Pt(18)
title1.paragraph_format.space_after = Pt(2)
add_run(title1, "ĐƠN ĐĂNG KÝ THAM GIA", size=14, bold=True)
title2 = doc.add_paragraph()
title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
title2.paragraph_format.space_after = Pt(14)
add_run(title2, "HỆ THỐNG MẠNG ĐẤU THẦU QUỐC GIA", size=14, bold=True)

add_rich_paragraph(doc, "Kính gửi: Trung tâm đấu thầu qua mạng quốc gia - Cục Quản lý Đấu thầu",
                    space_after=10)

# ---- Bang thong tin chinh (mo phong bieu mau goc) ----
LABEL_W = 6.5
VALUE_W = 9.5

rows_plain = [
    ("Vai trò:", "Nhà thầu"),
    ("Mã định danh:", "vn0111376819"),
    ("Mã hồ sơ:", ""),
]

section_1 = [
    ("Tên đơn vị (Tiếng Việt)", "Công ty TNHH Giải pháp Chuyển đổi số NEXUS"),
    ("Tên đơn vị (Tiếng Anh)", "NEXUS DIGITAL TRANSFORMATION SOLUTIONS COMPANY LIMITED"),
]

section_2 = [
    ("Mã số thuế", "0111376819"),
    ("Số GCN đăng ký kinh doanh/ Số GCN đầu tư", ""),
    ("Ngày cấp", "03/02/2026"),
    ("Quốc gia cấp", "Việt Nam"),
    ("Mã đơn vị quan hệ ngân sách", ""),
]

section_4 = [
    ("Địa chỉ", "Số nhà 256 đường Giải Phóng, Phường Phương Liệt, Thành phố Hà Nội"),
    ("Web", "https://nexuschuyendoiso.vn"),
    ("Mã vùng", ""),
]

section_5 = [
    ("Họ và tên", "Lê Thị Nguyệt Hằng"),
    ("Chức danh", "Giám đốc"),
    ("Phòng ban", ""),
    ("Loại chứng thực", "Căn cước công dân"),
    ("Số chứng thực", "001300006353"),
    ("Ngày cấp", "14/07/2025"),
    ("Quốc tịch", "Việt Nam"),
    ("Điện thoại", "0843256256"),
    ("Email", "Nexuschuyendoiso@gmail.com"),
]

section_6 = [
    ("Họ và tên", "Lê Thị Nguyệt Hằng"),
    ("Chức danh", "Giám đốc"),
    ("Phòng ban", "Giám đốc"),
    ("Loại chứng thực", "Căn cước công dân"),
    ("Quốc gia cấp", "Việt Nam"),
    ("Số chứng thực", "001300006353"),
    ("Ngày cấp", "14/07/2025"),
    ("Điện thoại", "0843256256"),
    ("Địa chỉ", "Số nhà 256 đường Giải Phóng, Phường Phương Liệt, Thành phố Hà Nội"),
    ("Email nhận hóa đơn điện tử", "ketoan.nexuschuyendoiso@gmail.com"),
    ("Email nhận thông tin chung", "Nexuschuyendoiso@gmail.com"),
]

section_2b = [
    ("Mã đơn vị quan hệ ngân sách", ""),
    ("Mã ngành", ""),
    ("Tên ngành", ""),
    ("Ngành chính", ""),
]

# Danh sach cac dong dang build: ("header", text) | ("plain", text) | ("pair", label, value)
plan = [("pair", l, v) for l, v in rows_plain]
plan += [("header", "I. Thông tin chung")]
plan += [("header", "1. Tên đơn vị")]
plan += [("pair", l, v) for l, v in section_1]
plan += [("header", "2. Thông tin chứng thực")]
plan += [("pair", l, v) for l, v in section_2]
plan += [("header", "3. Loại hình pháp lý")]
plan += [("plain", "Công ty trách nhiệm hữu hạn hai thành viên trở lên")]
plan += [("header", "4. Địa chỉ trụ sở")]
plan += [("pair", l, v) for l, v in section_4]
plan += [("header", "5. Thông tin người đại diện pháp luật của cơ quan, đơn vị")]
plan += [("pair", l, v) for l, v in section_5]
plan += [("header", "6. Thông tin người nhận thông báo")]
plan += [("pair", l, v) for l, v in section_6]
plan += [("header", "II. Thông tin dành cho Nhà thầu")]
plan += [("pair", l, v) for l, v in section_2b]

table = doc.add_table(rows=len(plan), cols=2)
set_table_borders(table, size=4, color='000000')
set_col_widths(table, [LABEL_W, VALUE_W])
table.alignment = WD_TABLE_ALIGNMENT.CENTER

for idx, item in enumerate(plan):
    cells = table.rows[idx].cells
    set_cell_no_margins(cells[0])
    set_cell_no_margins(cells[1])
    cells[0].text = ""
    cells[1].text = ""
    kind = item[0]
    if kind == "header":
        merged = cells[0].merge(cells[1])
        p = merged.paragraphs[0]
        add_run(p, item[1], size=13, bold=True)
    elif kind == "plain":
        merged = cells[0].merge(cells[1])
        p = merged.paragraphs[0]
        p.paragraph_format.line_spacing = 1.3
        add_run(p, item[1], size=13)
    else:
        _, label, value = item
        p0 = cells[0].paragraphs[0]
        p0.paragraph_format.line_spacing = 1.3
        add_run(p0, label, size=13)
        p1 = cells[1].paragraphs[0]
        p1.paragraph_format.line_spacing = 1.3
        add_run(p1, value, size=13)

# ---- Khoi ky ----
sig_table = doc.add_table(rows=1, cols=2)
remove_table_borders(sig_table)
set_col_widths(sig_table, [8.0, 8.0])
left, right = sig_table.rows[0].cells
set_cell_no_margins(left)
set_cell_no_margins(right)

p = left.paragraphs[0]
p.paragraph_format.space_before = Pt(10)
add_run(p, "Nơi nhận:", size=11, bold=True, italic=True)
for line in ["- Như trên;", "- Lưu văn thư."]:
    pn = left.add_paragraph()
    add_run(pn, line, size=11)

rp = right.paragraphs[0]
rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
rp.paragraph_format.space_before = Pt(10)
add_run(rp, "NGƯỜI ĐẠI DIỆN PHÁP LUẬT", size=13, bold=True)
for _ in range(4):
    right.add_paragraph()
rp2 = right.add_paragraph()
rp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(rp2, "Lê Thị Nguyệt Hằng", size=13, bold=True)

out_path = "/Users/hung/Projects/private/tieu_me/workspace/bao_cao/20260723_don_dang_ky_NEXUS_dau_thau.docx"
doc.save(out_path)
print("SAVED:", out_path)
