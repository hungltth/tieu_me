# -*- coding: utf-8 -*-
from docx import Document

W14 = '{http://schemas.microsoft.com/office/word/2010/wordml}'

SRC = "/Users/hung/Projects/private/tieu_me/workspace/bao_cao/20260716_phieutrinh_PGD_hoinghi_truc_tuyen_TT12.docx"
OUT = "/Users/hung/Projects/private/tieu_me/workspace/bao_cao/20260716_phieutrinh_PGD_hoinghi_truc_tuyen_TT12_v2.docx"

doc = Document(SRC)


def all_paragraphs(doc):
    paras = []

    def walk_body(container):
        for p in container.paragraphs:
            paras.append(p)
        for t in container.tables:
            for row in t.rows:
                for cell in row.cells:
                    walk_body(cell)

    walk_body(doc)
    return paras


PARAS = all_paragraphs(doc)
BY_ID = {}
for p in PARAS:
    pid = p._p.get(W14 + 'paraId')
    if pid:
        BY_ID[pid] = p

print("Indexed paragraphs:", len(BY_ID))


def replace_whole(paraid, new_text):
    p = BY_ID[paraid]
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    p.add_run(new_text)


# ---- Row 1: Tom tat noi dung cong viec (cap nhat theo Ke hoach moi 06/8/2026) ----
replace_whole('6F398B34',
    "        Thực hiện Thông tư số 12/2026/TT-BTC ngày 10/02/2026 của Bộ Tài chính; để phối hợp, "
    "hỗ trợ các cơ sở khám bệnh, chữa bệnh (KCB) BHYT chuẩn hóa 06 bảng danh mục sử dụng trong "
    "liên thông dữ liệu và thao tác lập, gửi biểu mẫu đề nghị thanh toán chi phí KCB BHYT trực "
    "tiếp trên Hệ thống thông tin giám định BHYT, Trung tâm đề xuất tổ chức Hội nghị trực tuyến "
    "hướng dẫn chuẩn hóa danh mục liên thông dữ liệu và lập biểu mẫu đề nghị thanh toán chi phí "
    "KCB BHYT theo Thông tư số 12/2026/TT-BTC. Cụ thể như sau:")

replace_whole('07914422',
    "+ Các đơn vị: Ban Thực hiện chính sách BHYT, Ban Tài chính - Kế toán, Ban Kiểm tra, Ban Kiểm "
    "toán nội bộ, Trung tâm Kiểm soát thanh toán BHXH, BHYT điện tử, Trung tâm Công nghệ thông tin "
    "và Chuyển đổi số")

replace_whole('2D845C9D',
    "b) Thời gian: 01 ngày, ngày 06/8/2026 (Thứ Năm); buổi sáng từ 8 giờ 00 phút (nội dung liên "
    "quan đến cơ sở KCB BHYT), buổi chiều từ 13 giờ 30 phút (nội dung liên quan đến cơ quan BHXH)")

replace_whole('6BE22C13',
    "- Hội nghị tổ chức theo hình thức trực tuyến, kết nối điểm cầu Bảo hiểm xã hội Việt Nam với "
    "điểm cầu Bảo hiểm xã hội các tỉnh, thành phố và điểm cầu Bảo hiểm xã hội cơ sở")

replace_whole('6A9DB825',
    "- Đầu mối chuẩn bị đường truyền, kỹ thuật: Văn phòng Bảo hiểm xã hội Việt Nam phối hợp Trung "
    "tâm Công nghệ thông tin và Chuyển đổi số")

replace_whole('78FDBEAC',
    "Trung tâm Kiểm soát thanh toán BHXH, BHYT điện tử chủ trì, phối hợp các đơn vị xây dựng nội "
    "dung, tài liệu hướng dẫn trình bày tại Hội nghị, gửi Phó Giám đốc Nguyễn Đức Hòa duyệt trước "
    "ngày 01/8/2026")

replace_whole('2661050F',
    "+ Trình bày hướng dẫn chuẩn hóa 06 bảng danh mục sử dụng trong liên thông dữ liệu KCB BHYT; "
    "hướng dẫn lập, duyệt Bảng kê chi tiết điều chỉnh (Mẫu 09/BH) và lập, đối chiếu số liệu các "
    "mẫu biểu đề nghị thanh toán (Mẫu 01/BH, 02/BH, 04/BH, 06/BH) theo Thông tư số 12/2026/TT-BTC")

replace_whole('4DC82714',
    "- Trung tâm Công nghệ thông tin và Chuyển đổi số: Cử cán bộ tham gia hướng dẫn, giải đáp các "
    "nội dung liên quan; phối hợp Văn phòng bảo đảm hạ tầng kỹ thuật, đường truyền phục vụ Hội "
    "nghị; mời đơn vị tư vấn công nghệ thông tin tham dự, phối hợp giải đáp vướng mắc")

replace_whole('57C3F14F',
    "- Văn phòng Bảo hiểm xã hội Việt Nam: Gửi văn bản thông báo thời gian, hình thức tổ chức Hội "
    "nghị; chuẩn bị hội trường, maket Hội nghị; chủ trì thiết lập kết nối và kiểm tra kết nối các "
    "điểm cầu (lần 1 ngày 05/8/2026, lần 2 ngày 06/8/2026); bảo đảm âm thanh, hình ảnh, đường "
    "truyền thông suốt trong suốt thời gian tổ chức Hội nghị")

doc.save(OUT)
print("SAVED:", OUT)
