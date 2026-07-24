# -*- coding: utf-8 -*-
from docx import Document

W14 = '{http://schemas.microsoft.com/office/word/2010/wordml}'

SRC = "/Users/hung/Projects/private/tieu_me/workspace/bao_cao/_tmp_phieutrinh_template.docx"
OUT = "/Users/hung/Projects/private/tieu_me/workspace/bao_cao/20260716_phieutrinh_PGD_hoinghi_truc_tuyen_TT12.docx"

doc = Document(SRC)


def all_paragraphs(doc):
    paras = []

    def walk_body(container):
        for p in container.paragraphs:
            paras.append(p)
        for t in container.tables:
            for row in t.rows:
                for cell in row.cells:
                    walk_body(cell)

    walk_body(doc)
    return paras


PARAS = all_paragraphs(doc)
BY_ID = {}
for p in PARAS:
    pid = p._p.get(W14 + 'paraId')
    if pid:
        BY_ID[pid] = p

print("Indexed paragraphs:", len(BY_ID))


def replace_whole(paraid, new_text):
    p = BY_ID[paraid]
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    p.add_run(new_text)


def replace_keep_label(paraid, new_text_after_label):
    p = BY_ID[paraid]
    runs = p.runs
    first = runs[0]
    for r in list(runs[1:]):
        r._element.getparent().remove(r._element)
    p.add_run(new_text_after_label)


# ---- Header section ----
replace_whole('6DF13A3F',
    "Vấn đề trình: Tổ chức Hội nghị trực tuyến hướng dẫn chuẩn hóa danh mục liên thông dữ liệu và "
    "lập biểu mẫu đề nghị thanh toán chi phí khám bệnh, chữa bệnh bảo hiểm y tế theo Thông tư số "
    "12/2026/TT-BTC ngày 10/02/2026 của Bộ Tài chính")

# ---- Row 1: Tom tat noi dung cong viec ----
replace_whole('6F398B34',
    "        Thực hiện Thông tư số 12/2026/TT-BTC ngày 10/02/2026 của Bộ Tài chính và trên cơ sở "
    "Kế hoạch số 668/KH-BHXH ngày 13/3/2026 của Bảo hiểm xã hội Việt Nam về tổ chức Hội nghị tập "
    "huấn trình tự, thủ tục giám định chi phí KCB BHYT; qua theo dõi, tổng hợp cho thấy các cơ sở "
    "khám bệnh, chữa bệnh (KCB) BHYT còn gặp một số khó khăn, vướng mắc trong việc chuẩn hóa 06 "
    "bảng danh mục sử dụng trong liên thông dữ liệu và thao tác lập, gửi dữ liệu đề nghị thanh toán "
    "chi phí KCB BHYT trên Hệ thống thông tin giám định BHYT. Để kịp thời hướng dẫn, tháo gỡ vướng "
    "mắc cho các cơ sở KCB BHYT, Trung tâm đề xuất tổ chức Hội nghị trực tuyến hướng dẫn chuẩn hóa "
    "danh mục liên thông dữ liệu và lập biểu mẫu đề nghị thanh toán chi phí KCB BHYT theo Thông tư "
    "số 12/2026/TT-BTC. Cụ thể như sau:")

replace_whole('07914422', "+ Các đơn vị: TTKS, CSYT, CNTT và Văn phòng BHXH Việt Nam")

replace_whole('2D845C9D',
    "b) Thời gian: 1/2 ngày (buổi sáng), từ 08 giờ 00 phút đến 12 giờ 00 phút, Thứ Sáu, ngày 31/7/2026")

replace_whole('7C45D7E2', "c) Hình thức")

replace_whole('6BE22C13',
    "- Hội nghị tổ chức theo hình thức trực tuyến, kết nối điểm cầu Bảo hiểm xã hội Việt Nam với "
    "điểm cầu Bảo hiểm xã hội các tỉnh, thành phố và các cơ sở khám bệnh, chữa bệnh BHYT")

replace_whole('6A9DB825',
    "- Đầu mối chuẩn bị đường truyền, kỹ thuật: Trung tâm Công nghệ thông tin")

# p[2FC74F41] "(chương trình chi tiết gửi kèm theo Phiếu trình này)" giữ nguyên

replace_whole('78FDBEAC',
    "Trung tâm Kiểm soát thanh toán BHXH, BHYT điện tử và Ban Thực hiện chính sách BHYT chuẩn bị "
    "nội dung, tài liệu hướng dẫn, gửi Phó Giám đốc Nguyễn Đức Hòa duyệt trước ngày 28/7/2026")

replace_whole('2661050F',
    "+ Trình bày hướng dẫn chuẩn hóa 06 bảng danh mục sử dụng trong liên thông dữ liệu KCB BHYT, "
    "lập Bảng kê chi tiết điều chỉnh (Mẫu 09/BH) và đối chiếu số liệu đề nghị thanh toán (Mẫu "
    "01/BH) theo Thông tư số 12/2026/TT-BTC ngày 10/02/2026")

replace_whole('510C0A80',
    "- Ban Thực hiện chính sách BHYT: Trình bày một số nội dung của Thông tư số 12/2026/TT-BTC "
    "liên quan đến cơ sở KCB BHYT và tham gia giải đáp vướng mắc")

replace_whole('4DC82714',
    "- Trung tâm Công nghệ thông tin: Bảo đảm hạ tầng kỹ thuật, đường truyền trực tuyến thông suốt "
    "trong suốt thời gian tổ chức Hội nghị")

replace_whole('57C3F14F',
    "- Văn phòng Bảo hiểm xã hội Việt Nam: Gửi Giấy mời, văn bản thông báo thời gian, hình thức tổ "
    "chức Hội nghị đến các đơn vị, Bảo hiểm xã hội các tỉnh, thành phố")

# ---- Row 3: de xuat cua Trung tam ----
replace_whole('3003190E',
    "Trung tâm dự thảo Kế hoạch tổ chức Hội nghị trực tuyến hướng dẫn chuẩn hóa danh mục liên "
    "thông dữ liệu và lập biểu mẫu đề nghị thanh toán chi phí KCB BHYT theo Thông tư số "
    "12/2026/TT-BTC")

doc.save(OUT)
print("SAVED:", OUT)
