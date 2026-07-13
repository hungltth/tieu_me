#!/usr/bin/env python3
"""
Tạo Báo cáo thành tích v3 từ file gốc:
 - Giữ nguyên toàn bộ nội dung gốc
 - Đỏ = điều chỉnh nội dung
 - Vàng (highlight) = bổ sung mới
 - Xanh + gạch ngang = đề xuất xóa
 - Bảng biểu cho số liệu đa tuyến, chuyên đề, sáng kiến
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ORIG = '/Users/hung/.claude/channels/telegram/inbox/1782960141740-AgAD3CEAAuVFMVY.docx'
OUT  = '/Users/hung/Projects/private/tieu_me/workspace/bao_cao/20260702_BCTICH_BangKhen_ThuTuong_THL_v3.docx'

RED   = (192,0,0)
BLUE  = (31,73,125)
BLACK = (0,0,0)

orig = Document(ORIG)
OP   = orig.paragraphs   # shorthand

doc = Document()
sec = doc.sections[0]
sec.page_height = Cm(29.7)
sec.page_width  = Cm(21.0)
sec.left_margin  = Cm(3.0)
sec.right_margin = Cm(2.0)
sec.top_margin   = Cm(2.5)
sec.bottom_margin = Cm(2.5)

# ── Helpers ──────────────────────────────────────────────────────────────────
def _rpr_font(run):
    rPr = run._element.get_or_add_rPr()
    rf  = rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts')
        rPr.append(rf)
    rf.set(qn('w:eastAsia'), 'Times New Roman')
    return rPr

def sf(run, size=13, bold=False, italic=False, color=None, hl=None, strike=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold   = bold
    run.italic = italic
    if strike: run.font.strike = True
    rPr = _rpr_font(run)
    if color: run.font.color.rgb = RGBColor(*color)
    if hl:
        h = OxmlElement('w:highlight')
        h.set(qn('w:val'), hl)
        rPr.append(h)

def ap(doc, text, bold=False, italic=False, color=None, hl=None, strike=False,
       align=None, fi=None, li=None, size=13, sa=4, sb=0):
    p  = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after  = Pt(sa)
    pf.space_before = Pt(sb)
    if fi is not None: pf.first_line_indent = Cm(fi)
    if li is not None: pf.left_indent = Cm(li)
    if align == 'center': p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right': p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == 'left':  p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:                  p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    sf(run, size=size, bold=bold, italic=italic, color=color, hl=hl, strike=strike)
    return p

def mp(doc, *parts, fi=None, li=None, sa=4, sb=0, align='justify', size=13):
    p  = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after  = Pt(sa)
    pf.space_before = Pt(sb)
    if fi is not None: pf.first_line_indent = Cm(fi)
    if li is not None: pf.left_indent = Cm(li)
    if align == 'center': p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:                 p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for pt in parts:
        txt    = pt[0]
        bold   = pt[1] if len(pt) > 1 else False
        hl     = pt[2] if len(pt) > 2 else None
        clr    = pt[3] if len(pt) > 3 else None
        strike = pt[4] if len(pt) > 4 else False
        run = p.add_run(txt)
        sf(run, size=size, bold=bold, color=clr, hl=hl, strike=strike)
    return p

def copy_para(doc, para):
    """Copy original paragraph preserving its text and basic format."""
    text = para.text
    if not text.strip():
        ap(doc, '', sa=2)
        return
    # Detect formatting from original
    bold = any(r.bold for r in para.runs) if para.runs else False
    size = 13
    for r in para.runs:
        if r.font.size:
            size = max(12, int(r.font.size.pt))
            break
    # Alignment
    if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
        align = 'center'
    else:
        align = 'justify'
    # Indentation guesses
    pf = para.paragraph_format
    fi = None
    li = None
    if pf.first_line_indent and pf.first_line_indent > 0:
        fi = round(pf.first_line_indent.cm, 2)
    if pf.left_indent and pf.left_indent > 0:
        li = round(pf.left_indent.cm, 2)
    # Heuristics for bullet lines
    if text.lstrip().startswith('-') or text.lstrip().startswith('+'):
        li = li or 0.5
        fi = None
    elif text.lstrip().startswith('*'):
        bold = True
        fi = None
    elif text.lstrip()[:3] in ('1. ','2. ','3. ','4. '):
        bold = True
    elif text.lstrip()[:3] in ('a) ','b) ','c) ','d) ','đ) ','e) '):
        bold = True
    ap(doc, text, bold=bold, size=size, align=align, fi=fi, li=li, sa=4)

def add_table_yellow(doc, headers, rows, col_widths=None):
    """Add a styled table with yellow header - labeled as new addition."""
    # Caption note
    ap(doc, '(Bảng tổng hợp – bổ sung mới)', italic=True, hl='yellow', size=11, sa=2)
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header
    hr = table.rows[0].cells
    for i, h in enumerate(headers):
        hr[i].text = h
        for run in hr[i].paragraphs[0].runs:
            sf(run, bold=True, size=11, hl='yellow')
        hr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data
    for ri, row in enumerate(rows):
        cells = table.rows[ri+1].cells
        is_total = row[0].startswith('Tổng') or row[0].startswith('Cộng')
        for ci, val in enumerate(row):
            cells[ci].text = val
            for run in cells[ci].paragraphs[0].runs:
                sf(run, size=11, bold=is_total, hl='yellow' if is_total else None)
            if ci > 0:
                cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)
    ap(doc, '', sa=2)

# ── TABLE DATA ────────────────────────────────────────────────────────────────
HDR_DT  = ['Năm', 'Số lượt (triệu)', 'Số tiền (tỷ đồng)']
ROWS_DT = [
    ('2017', '10,9',  '22.593,27'),
    ('2018', '13,3',  '23.481,49'),
    ('2019', '15,14', '26.195,28'),
    ('2020', '13,59', '25.246,97'),
    ('2021', '8,94',  '19.770,56'),
    ('2022', '12,75', '27.534,89'),
    ('2023', '15,76', '31.766,22'),
    ('2024', '17,12', '44.466,49'),
    ('2025', '18,7',  '51.452,97'),
    ('Tổng cộng', '127,44 triệu lượt', '272.508 tỷ đồng'),
]
CW_DT = [2.5, 4.5, 6.0]

HDR_CD  = ['Năm', 'Số CĐ (CĐ mới)', 'Tiền YC kiểm tra (tỷ đ)', 'Tiền từ chối/giảm trừ (tỷ đ)']
ROWS_CD = [
    ('2017', '15',           '583,8',   '115,1 (20,8%)'),
    ('2018', '46',           '415,1',   '123,1 (29,7%)'),
    ('2019', '121 (45 mới)', '192,9',   '74,1 (49,3%)'),
    ('2020', '~100 (22 mới)','—',        '14,75'),
    ('2021', '— (20 mới)',   '6,59',    '40,08'),
    ('2022', '94 (24 mới)',  '—',        '98,56'),
    ('2023', '87 (12 mới)',  '—',        '—'),
    ('2024', '90 (13 mới)',  '184,23',  '61,9'),
    ('Tổng 2017-2024', '402+ CĐ\n(146 CĐ mới)', '—', '~528 tỷ'),
]
CW_CD = [2.2, 3.5, 4.0, 4.3]

HDR_SK  = ['Năm', 'Sáng kiến / Đề tài KHCN', 'Cấp', 'QĐ công nhận']
ROWS_SK = [
    ('2018/2019', 'Liên thông quản lý chi KCB BHYT trên các phần mềm nghiệp vụ', 'Ngành', 'QĐ 233/QĐ-BHXH\n26/02/2019'),
    ('2020',      'Bản đồ theo dõi người mắc bệnh nền hỗ trợ phòng chống Covid-19', 'Cơ sở', 'QĐ 29/QĐ-GĐĐT\n29/12/2020'),
    ('2021',      'Công cụ xác định quỹ và giám sát chi KCB BHYT theo định suất', 'Cơ sở', 'QĐ 33/QĐ-GĐĐT\n08/11/2021'),
    ('2022',      'Cấu hình nhóm tỉnh trên phần mềm Giám sát', 'Cơ sở', 'QĐ 27/QĐ-GĐĐT\n09/11/2022'),
    ('2022',      'Hoàn thiện quy trình giám định BHYT (đồng tác giả)', 'Bộ', 'QĐ 1104/QĐ-BHXH\n22/4/2022'),
    ('2023',      'Công cụ lập báo cáo quyết toán đảm bảo tiến độ', 'Cơ sở', 'QĐ 20/QĐ-GĐĐT\n03/11/2023'),
    ('2024',      'Cảnh báo gia tăng chi phí KCB theo Nghị định 75/2023/NĐ-CP', 'Cơ sở', 'QĐ 24/QĐ-GĐĐT\n29/10/2024'),
    ('2023–2025', 'AI-OCR trong giám định BHYT – Chủ nhiệm đề tài KHCN cơ sở', 'Cơ sở', 'QĐ 12/QĐ-BHXH\n08/01/2025'),
    ('2025',      'Công cụ công bố thuốc, TTBYT chưa sử dụng hết theo kết quả trúng thầu', 'Cơ sở', 'QĐ 04/QĐ-TTKS\n04/02/2026'),
    ('2025',      'Chức năng thống kê KCB theo đơn vị tham gia BHXH trên PM Giám sát', 'Cơ sở', 'QĐ 04/QĐ-TTKS\n04/02/2026'),
]
CW_SK = [2.0, 7.0, 1.8, 3.2]

# ── MODIFIED TEXT ─────────────────────────────────────────────────────────────
MOD = {
    14: 'Quyền hạn, nhiệm vụ được giao hoặc đảm nhận từ năm 2016 đến tháng 5/2026',
    36: ('Từ việc chỉ đạo sát sao và tích cực hưởng ứng các phong trào thi đua, '
         'tôi đã đạt được nhiều thành tích nổi bật trong suốt quá trình công tác '
         'gắn liền với ứng dụng công nghệ thông tin và thúc đẩy chuyển đổi số '
         'trong ngành Bảo hiểm xã hội.'),
    39: ('Với vai trò là Phó Giám đốc trực tiếp phụ trách lĩnh vực công nghệ thông tin '
         'và chuyển đổi số, tôi đã chủ động tham mưu và chỉ đạo toàn diện các mảng '
         'công tác then chốt: quản lý, vận hành ổn định và không ngừng phát triển '
         'Hệ thống thông tin giám định BHYT; khai thác, phân tích dữ liệu, xây dựng '
         'các chuyên đề giám định; thực hiện thanh toán đa tuyến ngoại tỉnh; '
         'thúc đẩy chuyển đổi số và xây dựng cơ sở dữ liệu quốc gia về BHYT; '
         'đồng thời hoàn thành tốt các công tác công đoàn trong đơn vị:'),
    87: ('Chủ động tham mưu cho Thủ trưởng đơn vị, Lãnh đạo BHXH Việt Nam ban hành '
         'các văn bản hướng dẫn liên thông dữ liệu, chuẩn hoá danh mục, quản lý, '
         'khai thác thông tin trên Hệ thống; chỉ đạo BHXH các tỉnh, thành phố thực '
         'hiện đầy đủ quy trình giám định điện tử, quy định về liên thông các phần '
         'mềm nghiệp vụ để quản lý công tác chi khám chữa bệnh BHYT và thực hiện '
         'quyết toán trên Hệ thống; thẩm định quyết toán tập trung từ năm 2020, '
         'đảm bảo sự thống nhất về thời điểm, số liệu, tiến độ quyết toán; tham mưu '
         'phối hợp với các đơn vị xây dựng các văn bản hướng dẫn, chủ động chỉ đạo '
         'phân tích dữ liệu, hỗ trợ BHXH các tỉnh, thành phố xác định tổng mức '
         'thanh toán, lập báo cáo thẩm định trên phần mềm.'),
}

# Paragraph [38] đề xuất chuyển về mục d)
DELETE_PARAS = {38}

# ── NEW ADDITIONS (yellow) after specific paragraph indices ──────────────────
def ins_after_13(doc):
    ap(doc,
       'Trong suốt quá trình công tác từ năm 2016 đến tháng 5/2026, với vai trò '
       'là chuyên gia hàng đầu về ứng dụng công nghệ thông tin và chuyển đổi số '
       'trong lĩnh vực Bảo hiểm y tế, bản thân tôi đã có nhiều đóng góp quan trọng '
       'trong việc xây dựng, vận hành và phát triển Hệ thống thông tin giám định BHYT '
       '— một trong những hệ thống CNTT trọng điểm của Ngành, tiếp nhận và xử lý '
       'hàng trăm triệu lượt khám chữa bệnh mỗi năm, góp phần bảo vệ quyền lợi '
       'cho hàng chục triệu người tham gia BHYT và bảo vệ tính bền vững của quỹ BHYT.',
       hl='yellow', fi=1.27, sa=4)

def ins_after_59(doc):
    ap(doc,
       'Bản thân tôi đã trực tiếp tham gia nghiên cứu, đề xuất và thẩm định '
       'nhiều quy định kỹ thuật quan trọng trong Thông tư số 12/2026/TT-BTC, '
       'đặc biệt các quy định về quy trình giám định điện tử, chuẩn hóa biểu mẫu '
       'tổng hợp thanh toán và tích hợp với cơ sở dữ liệu quốc gia, từ đó đảm bảo '
       'Hệ thống được nâng cấp đồng bộ và triển khai kịp thời khi Thông tư có hiệu '
       'lực thi hành trên toàn hệ thống BHXH Việt Nam.',
       hl='yellow', fi=1.27, sa=4)

def ins_after_71(doc):
    ap(doc,
       'Tổng hợp kết quả xây dựng và thực hiện chuyên đề giám định BHYT '
       '(2017 – 2024):',
       bold=True, hl='yellow', fi=1.27, sa=2)
    add_table_yellow(doc, HDR_CD, ROWS_CD, CW_CD)

def ins_after_85(doc):
    ap(doc,
       'Tổng hợp kết quả thanh toán đa tuyến ngoại tỉnh (2017 – 2025):',
       bold=True, hl='yellow', fi=1.27, sa=2)
    add_table_yellow(doc, HDR_DT, ROWS_DT, CW_DT)
    ap(doc,
       'Lũy kế 09 năm (2017–2025), Trung tâm đã tổng hợp, thông báo thanh toán '
       'đa tuyến ngoại tỉnh cho tổng cộng hơn 127 triệu lượt khám chữa bệnh với '
       'tổng số tiền hơn 272.500 tỷ đồng, đảm bảo quyền lợi khám chữa bệnh '
       'kịp thời cho người tham gia BHYT trên toàn quốc.',
       hl='yellow', fi=1.27, sa=4)

def ins_after_89(doc):
    ap(doc,
       'Hệ thống thông tin giám định BHYT — công trình mà tôi tham gia xây dựng '
       'và phát triển từ những ngày đầu — đã được Hiệp hội An sinh xã hội quốc tế '
       '(ISSA) và Hiệp hội An sinh xã hội châu Á – Thái Bình Dương (ASSA) '
       'trao tặng giải thưởng về ứng dụng công nghệ thông tin xuất sắc trong '
       'quản lý an sinh xã hội. Đây là sự ghi nhận của cộng đồng quốc tế đối với '
       'những thành tựu tiêu biểu mà Việt Nam đã đạt được trong lĩnh vực chuyển '
       'đổi số và ứng dụng CNTT nâng cao hiệu quả giám định, thanh toán BHYT.',
       hl='yellow', fi=1.27, sa=4)

def ins_after_99(doc):
    ap(doc,
       'Định hướng phát triển tiếp theo mà tôi đang tập trung tham mưu gồm: '
       'ứng dụng trí tuệ nhân tạo (AI) phân tích dự báo nhu cầu khám chữa bệnh; '
       'hoàn thiện hệ thống AI-OCR để tự động hóa đối soát danh mục thuốc, '
       'dịch vụ kỹ thuật y tế; xây dựng nền tảng dữ liệu tích hợp phục vụ quản '
       'lý quỹ BHYT thông minh và phát hiện gian lận tự động, đáp ứng yêu cầu '
       'của Luật BHYT số 51/2024/QH15 và Nghị định số 188/2025/NĐ-CP.',
       hl='yellow', fi=1.27, sa=4)

def ins_after_156(doc):
    ap(doc,
       'Tổng hợp sáng kiến, đề tài KHCN đã được công nhận (2018 – 2025):',
       bold=True, hl='yellow', fi=1.27, sa=2)
    add_table_yellow(doc, HDR_SK, ROWS_SK, CW_SK)
    ap(doc,
       'Tổng cộng: 10 sáng kiến/đề tài KHCN được công nhận (01 cấp Ngành, '
       '01 cấp Bộ, 08 cấp cơ sở), phản ánh quá trình nghiên cứu, sáng tạo '
       'liên tục và ứng dụng thực tiễn cao trong giai đoạn 2018–2026.',
       hl='yellow', fi=1.27, sa=4)

AFTER_HOOKS = {
    13:  ins_after_13,
    59:  ins_after_59,
    71:  ins_after_71,
    85:  ins_after_85,
    89:  ins_after_89,
    99:  ins_after_99,
    156: ins_after_156,
}

# ── MAIN LOOP ─────────────────────────────────────────────────────────────────
for i, para in enumerate(OP):
    text = para.text

    if i in MOD:
        # Show MODIFIED version in red
        new_txt = MOD[i]
        # Determine formatting from original
        bold  = any(r.bold for r in para.runs) if para.runs else False
        size  = 13
        for r in para.runs:
            if r.font.size: size = max(12, int(r.font.size.pt)); break
        fi = None; li = None
        if para.paragraph_format.first_line_indent and para.paragraph_format.first_line_indent > 0:
            fi = round(para.paragraph_format.first_line_indent.cm, 2)
        if para.paragraph_format.left_indent and para.paragraph_format.left_indent > 0:
            li = round(para.paragraph_format.left_indent.cm, 2)
        align = 'center' if para.alignment == WD_ALIGN_PARAGRAPH.CENTER else 'justify'
        ap(doc, new_txt, bold=bold, color=RED, size=size, align=align, fi=fi, li=li, sa=4)

    elif i in DELETE_PARAS:
        # Show original with strikethrough + blue as "đề xuất xóa"
        if text.strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.first_line_indent = Cm(1.27)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # Note prefix
            note_run = p.add_run('[Đề xuất chuyển xuống mục d)]: ')
            sf(note_run, size=11, bold=True, color=BLUE, italic=True)
            # Original text with strikethrough
            content_run = p.add_run(text)
            sf(content_run, size=13, color=BLUE, strike=True)

    else:
        # Keep original unchanged
        copy_para(doc, para)

    # Insert after this paragraph
    if i in AFTER_HOOKS:
        AFTER_HOOKS[i](doc)

# ── LEGEND NOTE ───────────────────────────────────────────────────────────────
ap(doc, '', sa=2)
ap(doc, '─' * 60, align='center', size=10, sa=2)
mp(doc,
   ('CHÚ THÍCH MÀU SẮC: ', True, None, None),
   ('Chữ đỏ', False, None, RED),
   (' = điều chỉnh nội dung | ', False, None, None),
   ('Nền vàng', False, 'yellow', None),
   (' = bổ sung mới | ', False, None, None),
   ('Chữ xanh gạch ngang', False, None, BLUE, True),
   (' = đề xuất xóa/chuyển mục', False, None, None),
   size=11, align='center', sa=2)

doc.save(OUT)
print(f'Saved: {OUT}')
