# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Times New Roman"
RED = (192, 0, 0)

YELLOW_TERMS = [
    "Thông tư số 12/2026/TT-BTC",
    "Nghị định số 188/2025/NĐ-CP",
    "Kế hoạch số 668/KH-BHXH",
    "Mẫu 09/BH",
    "Mẫu 01/BH",
    "06 bảng danh mục",
]
CYAN_TERMS = [
    "Trung tâm Kiểm soát thanh toán BHXH, BHYT điện tử",
    "Ban Thực hiện chính sách BHYT",
    "ngày 31/7/2026",
    "hình thức trực tuyến",
]
ALL_TERMS = sorted(set(YELLOW_TERMS + CYAN_TERMS), key=len, reverse=True)


RPR_ORDER = ['rStyle', 'rFonts', 'b', 'bCs', 'i', 'iCs', 'caps', 'smallCaps', 'strike',
             'dstrike', 'outline', 'shadow', 'emboss', 'imprint', 'noProof', 'snapToGrid',
             'vanish', 'webHidden', 'color', 'spacing', 'w', 'kern', 'position', 'sz', 'szCs',
             'highlight', 'u', 'effect', 'bdr', 'shd', 'fitText', 'vertAlign', 'rtl', 'cs',
             'em', 'lang', 'eastAsianLayout', 'specVanish', 'oMath']

TBLPR_ORDER = ['tblStyle', 'tblpPr', 'tblOverlap', 'bidiVisual', 'tblStyleRowBandSize',
               'tblStyleColBandSize', 'tblW', 'jc', 'tblCellSpacing', 'tblInd', 'tblBorders',
               'shd', 'tblLayout', 'tblCellMar', 'tblLook', 'tblCaption', 'tblDescription']


def insert_in_order(parent, new_el, order_list):
    tag = new_el.tag.split('}')[-1]
    try:
        new_idx = order_list.index(tag)
    except ValueError:
        parent.append(new_el)
        return
    for child in parent:
        child_tag = child.tag.split('}')[-1]
        if child_tag in order_list and order_list.index(child_tag) > new_idx:
            child.addprevious(new_el)
            return
    parent.append(new_el)


def set_font(run, size=13, bold=False, italic=False, color=None, highlight=None, underline=False):
    run.font.name = FONT
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        insert_in_order(rPr, rFonts, RPR_ORDER)
    rFonts.set(qn('w:ascii'), FONT)
    rFonts.set(qn('w:hAnsi'), FONT)
    rFonts.set(qn('w:eastAsia'), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline
    if color:
        run.font.color.rgb = RGBColor(*color)
    if highlight:
        hl = OxmlElement('w:highlight')
        hl.set(qn('w:val'), highlight)
        insert_in_order(rPr, hl, RPR_ORDER)


def add_run(p, text, **kw):
    r = p.add_run(text)
    set_font(r, **kw)
    return r


def add_rich_paragraph(doc_or_cell, text, size=13, align=None, space_after=6, space_before=0,
                        first_line_indent=None, line_spacing=1.3):
    p = doc_or_cell.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = line_spacing
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)

    remaining = text
    segments = []
    while remaining:
        earliest_idx = None
        earliest_term = None
        for t in ALL_TERMS:
            idx = remaining.find(t)
            if idx != -1 and (earliest_idx is None or idx < earliest_idx):
                earliest_idx = idx
                earliest_term = t
        if earliest_idx is None:
            segments.append((remaining, None))
            break
        if earliest_idx > 0:
            segments.append((remaining[:earliest_idx], None))
        segments.append((earliest_term, 'yellow' if earliest_term in YELLOW_TERMS else 'cyan'))
        remaining = remaining[earliest_idx + len(earliest_term):]

    for seg_text, hl in segments:
        add_run(p, seg_text, size=size, highlight=hl)
    return p


def add_heading(doc, text, size=13, space_before=10, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.3
    add_run(p, text, size=size, bold=True, color=RED)
    return p


def set_cell_no_margins(cell, top=0, bottom=0, left=0, right=0):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement('w:tcMar')
    for tag, val in (('top', top), ('start', left), ('bottom', bottom), ('end', right)):
        el = OxmlElement(f'w:{tag}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        mar.append(el)
    tcPr.append(mar)


def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'nil')
        borders.append(el)
    insert_in_order(tblPr, borders, TBLPR_ORDER)


def set_table_borders(table, size=4, color='000000'):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(size))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        borders.append(el)
    insert_in_order(tblPr, borders, TBLPR_ORDER)


def shade_cell(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def set_col_widths(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for idx, w in enumerate(widths_cm):
            row.cells[idx].width = Cm(w)
    for idx, w in enumerate(widths_cm):
        table.columns[idx].width = Cm(w)


def new_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    style = doc.styles['Normal']
    style.font.name = FONT
    style.font.size = Pt(13)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT)
    return doc


def add_header_block(doc):
    table = doc.add_table(rows=1, cols=2)
    remove_table_borders(table)
    set_col_widths(table, [6.0, 10.0])
    left, right = table.rows[0].cells
    set_cell_no_margins(left)
    set_cell_no_margins(right)

    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "BỘ TÀI CHÍNH", size=13)
    p2 = left.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p2, "BẢO HIỂM XÃ HỘI VIỆT NAM", size=13, bold=True, underline=True)
    p3 = left.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(30)
    add_run(p3, "Số:      /KH-BHXH", size=13)

    q = right.paragraphs[0]
    q.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(q, "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", size=13, bold=True)
    q2 = right.add_paragraph()
    q2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(q2, "Độc lập - Tự do - Hạnh phúc", size=14, bold=True, underline=True)
    q3 = right.add_paragraph()
    q3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    q3.paragraph_format.space_before = Pt(30)
    add_run(q3, "Hà Nội, ngày     tháng     năm 2026", size=13, italic=True)
    return table


def add_title_block(doc, ten_loai_vb, trich_yeu_lines):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, ten_loai_vb, size=14, bold=True)
    for i, line in enumerate(trich_yeu_lines):
        pl = doc.add_paragraph()
        pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pl.paragraph_format.space_after = Pt(2) if i < len(trich_yeu_lines) - 1 else Pt(14)
        add_run(pl, line, size=14, bold=True)


# ============================================================
# 1. KE HOACH
# ============================================================
doc = new_doc()
add_header_block(doc)
add_title_block(doc, "KẾ HOẠCH", [
    "Tổ chức Hội nghị trực tuyến hướng dẫn chuẩn hóa danh mục liên thông dữ liệu",
    "và lập biểu mẫu đề nghị thanh toán chi phí khám bệnh, chữa bệnh bảo hiểm y tế",
    "theo Thông tư số 12/2026/TT-BTC ngày 10/02/2026 của Bộ Tài chính",
])

intro = (
    "Thực hiện Thông tư số 12/2026/TT-BTC ngày 10/02/2026 của Bộ Tài chính quy định "
    "trình tự, thủ tục giám định chi phí khám bệnh, chữa bệnh (KCB) bảo hiểm y tế (BHYT), "
    "biểu mẫu tổng hợp thanh toán, quyết toán và biện pháp thi hành Nghị định số 188/2025/NĐ-CP "
    "ngày 01/7/2025 của Chính phủ và Kế hoạch số 668/KH-BHXH ngày 13/3/2026 của Bảo hiểm xã hội "
    "Việt Nam về tổ chức Hội nghị tập huấn trình tự, thủ tục giám định chi phí KCB BHYT; để tiếp "
    "tục hỗ trợ các cơ sở khám bệnh, chữa bệnh (KCB) BHYT chuẩn hóa 06 bảng danh mục sử dụng trong "
    "liên thông dữ liệu và thao tác lập, gửi dữ liệu đề nghị thanh toán chi phí KCB BHYT trực tiếp "
    "trên phần mềm, Bảo hiểm xã hội Việt Nam ban hành Kế hoạch tổ chức Hội nghị trực tuyến hướng "
    "dẫn, cụ thể như sau:"
)
add_rich_paragraph(doc, intro, first_line_indent=1.25, space_after=6)

add_heading(doc, "I. MỤC ĐÍCH, YÊU CẦU")
muc_dich = [
    "1. Hướng dẫn các cơ sở KCB BHYT chuẩn hóa 06 bảng danh mục sử dụng trong liên thông dữ liệu "
    "KCB BHYT theo quy định tại Thông tư số 12/2026/TT-BTC ngày 10/02/2026 của Bộ Tài chính.",
    "2. Hướng dẫn cơ sở KCB BHYT thao tác lập Bảng kê chi tiết điều chỉnh (Mẫu 09/BH) và đối chiếu "
    "số liệu đề nghị thanh toán (Mẫu 01/BH) trực tiếp trên Hệ thống thông tin giám định BHYT.",
    "3. Kịp thời trao đổi, giải đáp vướng mắc phát sinh của cơ sở KCB BHYT trong quá trình chuẩn "
    "hóa danh mục và lập, gửi dữ liệu đề nghị thanh toán chi phí KCB BHYT.",
    "4. Hội nghị tổ chức đảm bảo thiết thực, hiệu quả, tiết kiệm.",
]
for line in muc_dich:
    add_rich_paragraph(doc, line, first_line_indent=1.25)

add_heading(doc, "II. NỘI DUNG HỘI NGHỊ")
add_rich_paragraph(doc, "Hội nghị tổ chức trong 1/2 ngày với các nội dung chính:", first_line_indent=1.25)
noi_dung = [
    "1. Một số nội dung của Thông tư số 12/2026/TT-BTC liên quan đến cơ sở KCB BHYT (30 phút).",
    "2. Hướng dẫn một số nội dung lập 06 bảng danh mục sử dụng trong liên thông dữ liệu KCB BHYT (60 phút).",
    "3. Hướng dẫn lập Bảng kê chi tiết điều chỉnh (Mẫu 09/BH) (30 phút).",
    "4. Hướng dẫn đối chiếu số liệu đề nghị thanh toán (Mẫu 01/BH) (15 phút).",
    "5. Giải lao (15 phút).",
    "6. Giải đáp vướng mắc của cơ sở KCB BHYT (60 phút).",
]
for line in noi_dung:
    add_rich_paragraph(doc, line, first_line_indent=1.25)
add_rich_paragraph(doc, "(Chương trình chi tiết tại Phụ lục kèm theo Kế hoạch này)", first_line_indent=1.25)

add_heading(doc, "III. THỜI GIAN, HÌNH THỨC VÀ THÀNH PHẦN THAM DỰ")
add_rich_paragraph(doc, "1. Thời gian: 1/2 ngày (buổi sáng), Thứ Sáu, ngày 31/7/2026.", first_line_indent=1.25)
add_rich_paragraph(doc,
    "2. Hình thức: Hội nghị tổ chức theo hình thức trực tuyến, kết nối điểm cầu Bảo hiểm xã hội "
    "Việt Nam với điểm cầu Bảo hiểm xã hội các tỉnh, thành phố và các cơ sở KCB BHYT.",
    first_line_indent=1.25)
add_rich_paragraph(doc, "3. Thành phần tham dự:", first_line_indent=1.25)
thanh_phan = [
    "- Điểm cầu Bảo hiểm xã hội Việt Nam: Lãnh đạo Ban Thực hiện chính sách BHYT; lãnh đạo và viên "
    "chức Trung tâm Kiểm soát thanh toán BHXH, BHYT điện tử; viên chức Trung tâm Công nghệ thông "
    "tin phụ trách kỹ thuật đường truyền.",
    "- Điểm cầu Bảo hiểm xã hội các tỉnh, thành phố: lãnh đạo phòng và viên chức làm công tác giám "
    "định chi phí KCB BHYT.",
    "- Điểm cầu cơ sở KCB BHYT: lãnh đạo phụ trách công tác BHYT và cán bộ, viên chức trực tiếp "
    "thực hiện việc chuẩn hóa danh mục, lập, gửi dữ liệu đề nghị thanh toán chi phí KCB BHYT.",
]
for line in thanh_phan:
    add_rich_paragraph(doc, line, first_line_indent=1.25)

add_heading(doc, "IV. KINH PHÍ THỰC HIỆN")
add_rich_paragraph(doc,
    "Kinh phí tổ chức Hội nghị từ nguồn chi quản lý bảo hiểm xã hội, bảo hiểm thất nghiệp, bảo "
    "hiểm y tế năm 2026.", first_line_indent=1.25)

add_heading(doc, "V. TỔ CHỨC THỰC HIỆN")
to_chuc = [
    "1. Trung tâm Kiểm soát thanh toán BHXH, BHYT điện tử: Chủ trì, phối hợp với Ban Thực hiện "
    "chính sách BHYT xây dựng nội dung, tài liệu hướng dẫn; chuẩn bị kịch bản, đường truyền trực "
    "tuyến; tổng hợp, giải đáp vướng mắc của cơ sở KCB BHYT tại Hội nghị.",
    "2. Ban Thực hiện chính sách BHYT: Phối hợp xây dựng nội dung liên quan đến Thông tư số "
    "12/2026/TT-BTC; cử cán bộ tham gia hướng dẫn, giải đáp tại Hội nghị.",
    "3. Trung tâm Công nghệ thông tin: Bảo đảm hạ tầng kỹ thuật, đường truyền trực tuyến thông "
    "suốt trong suốt thời gian tổ chức Hội nghị.",
    "4. Văn phòng Bảo hiểm xã hội Việt Nam: Gửi Giấy mời, văn bản thông báo thời gian, hình thức "
    "tổ chức Hội nghị đến các đơn vị, Bảo hiểm xã hội các tỉnh, thành phố.",
    "5. Bảo hiểm xã hội các tỉnh, thành phố: Thông báo, đôn đốc các cơ sở KCB BHYT trên địa bàn "
    "tham dự đầy đủ, đúng thành phần; tổng hợp các nội dung vướng mắc (nếu có) gửi về Bảo hiểm xã "
    "hội Việt Nam qua Trung tâm Kiểm soát thanh toán BHXH, BHYT điện tử trước 05 ngày tổ chức Hội nghị.",
    "6. Cơ sở KCB BHYT: Cử cán bộ, viên chức tham dự đầy đủ, đúng thành phần; nghiên cứu trước tài "
    "liệu hướng dẫn; chuẩn bị đường truyền, phương tiện kỹ thuật tại điểm cầu.",
]
for line in to_chuc:
    add_rich_paragraph(doc, line, first_line_indent=1.25)

ket = (
    "Trên đây là Kế hoạch tổ chức Hội nghị trực tuyến hướng dẫn chuẩn hóa danh mục liên thông dữ "
    "liệu và lập biểu mẫu đề nghị thanh toán chi phí KCB BHYT theo Thông tư số 12/2026/TT-BTC. Bảo "
    "hiểm xã hội Việt Nam yêu cầu các đơn vị thuộc và trực thuộc, Bảo hiểm xã hội các tỉnh, thành "
    "phố nghiêm túc triển khai thực hiện. Trong quá trình thực hiện nếu có khó khăn, vướng mắc, đề "
    "nghị phản ánh kịp thời về Bảo hiểm xã hội Việt Nam (qua Trung tâm Kiểm soát thanh toán BHXH, "
    "BHYT điện tử) để xem xét, giải quyết./."
)
add_rich_paragraph(doc, ket, first_line_indent=1.25, space_before=8)

sig_table = doc.add_table(rows=1, cols=2)
remove_table_borders(sig_table)
set_col_widths(sig_table, [9.0, 7.0])
left, right = sig_table.rows[0].cells
set_cell_no_margins(left)
set_cell_no_margins(right)

p = left.paragraphs[0]
p.paragraph_format.space_before = Pt(6)
add_run(p, "Nơi nhận:", size=11, bold=True, italic=True)
noi_nhan = [
    "- Giám đốc BHXH Việt Nam (để b/c);",
    "- Các Phó Giám đốc BHXH Việt Nam (để b/c);",
    "- Các đơn vị: Ban CSYT, TTKS, CNTT, VP (để t/h);",
    "- BHXH các tỉnh, thành phố (để t/h);",
    "- Lưu: VT, TTKS.",
]
for line in noi_nhan:
    pn = left.add_paragraph()
    add_run(pn, line, size=11)

rp = right.paragraphs[0]
rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(rp, "TL. GIÁM ĐỐC", size=13, bold=True)
rp2 = right.add_paragraph()
rp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(rp2, "GIÁM ĐỐC TRUNG TÂM KIỂM SOÁT THANH TOÁN BHXH, BHYT ĐIỆN TỬ", size=13, bold=True)
for _ in range(4):
    right.add_paragraph()
rp5 = right.add_paragraph()
rp5.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(rp5, "(Họ và tên)", size=11, italic=True)

# ============================================================
# 2. PHU LUC - CHUONG TRINH (new page, same file)
# ============================================================
doc.add_section(WD_SECTION.NEW_PAGE)
sec2 = doc.sections[-1]
sec2.page_width = Cm(21.0)
sec2.page_height = Cm(29.7)
sec2.top_margin = Cm(2.0)
sec2.bottom_margin = Cm(2.0)
sec2.left_margin = Cm(2.0)
sec2.right_margin = Cm(2.0)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.space_after = Pt(2)
add_run(p, "PHỤ LỤC", size=13, bold=True, italic=True)
add_rich_paragraph(doc, "(Kèm theo Kế hoạch số   /KH-BHXH ngày   tháng   năm 2026 của Bảo hiểm xã hội Việt Nam)",
                    size=12, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=16)

title1 = doc.add_paragraph()
title1.alignment = WD_ALIGN_PARAGRAPH.CENTER
title1.paragraph_format.space_after = Pt(2)
add_run(title1, "DỰ KIẾN CHƯƠNG TRÌNH", size=14, bold=True)
for line in [
    "HỘI NGHỊ TRỰC TUYẾN HƯỚNG DẪN CHUẨN HÓA DANH MỤC LIÊN THÔNG DỮ LIỆU",
    "VÀ LẬP BIỂU MẪU ĐỀ NGHỊ THANH TOÁN CHI PHÍ KCB BHYT",
    "THEO THÔNG TƯ SỐ 12/2026/TT-BTC NGÀY 10/02/2026 CỦA BỘ TÀI CHÍNH",
]:
    tl = doc.add_paragraph()
    tl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tl.paragraph_format.space_after = Pt(2)
    add_run(tl, line, size=14, bold=True)

tg = doc.add_paragraph()
tg.alignment = WD_ALIGN_PARAGRAPH.CENTER
tg.paragraph_format.space_before = Pt(6)
tg.paragraph_format.space_after = Pt(14)
add_run(tg, "Thời gian: Sáng Thứ Sáu, ngày 31 tháng 7 năm 2026", size=13, italic=True, bold=True)

rows = [
    ("8h00 - 8h10", "Đón tiếp đại biểu, tuyên bố lý do, giới thiệu đại biểu",
     "Đại diện lãnh đạo Trung tâm Kiểm soát thanh toán BHXH, BHYT điện tử"),
    ("8h10 - 8h20", "Phát biểu khai mạc Hội nghị", "Lãnh đạo Bảo hiểm xã hội Việt Nam"),
    ("8h20 - 8h50", "1. Một số nội dung của Thông tư số 12/2026/TT-BTC liên quan đến cơ sở KCB BHYT (30 phút)",
     "Ban Thực hiện chính sách BHYT"),
    ("8h50 - 9h50", "2. Hướng dẫn một số nội dung lập 06 bảng danh mục sử dụng trong liên thông dữ liệu KCB BHYT (60 phút)",
     "Trung tâm Kiểm soát thanh toán BHXH, BHYT điện tử"),
    ("9h50 - 10h20", "3. Hướng dẫn lập Bảng kê chi tiết điều chỉnh (Mẫu 09/BH) (30 phút)",
     "Trung tâm Kiểm soát thanh toán BHXH, BHYT điện tử"),
    ("10h20 - 10h35", "4. Hướng dẫn đối chiếu số liệu đề nghị thanh toán (Mẫu 01/BH) (15 phút)",
     "Trung tâm Kiểm soát thanh toán BHXH, BHYT điện tử"),
    ("10h35 - 10h50", "5. Giải lao (15 phút)", ""),
    ("10h50 - 11h50", "6. Giải đáp vướng mắc của cơ sở KCB BHYT (60 phút)",
     "Ban Thực hiện chính sách BHYT, Trung tâm Kiểm soát thanh toán BHXH, BHYT điện tử"),
    ("11h50 - 12h00", "Kết luận, bế mạc Hội nghị", "Lãnh đạo Bảo hiểm xã hội Việt Nam"),
]

table = doc.add_table(rows=len(rows) + 1, cols=3)
set_table_borders(table, size=4, color='000000')
set_col_widths(table, [3.2, 9.3, 4.5])
table.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr = table.rows[0].cells
headers = ["Thời gian", "Nội dung", "Thực hiện"]
for i, h in enumerate(headers):
    hdr[i].text = ""
    p = hdr[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, h, size=13, bold=True)
    shade_cell(hdr[i], 'D9D9D9')

for r_idx, (tg_val, nd_val, th_val) in enumerate(rows, start=1):
    cells = table.rows[r_idx].cells
    cells[0].text = ""
    p0 = cells[0].paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p0, tg_val, size=12, bold=True)

    cells[1].text = ""
    add_rich_paragraph(cells[1], nd_val, size=12, space_after=0)

    cells[2].text = ""
    p2c = cells[2].paragraphs[0]
    p2c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p2c, th_val, size=12)

note = doc.add_paragraph()
note.paragraph_format.space_before = Pt(10)
add_run(note, "(Lịch làm việc có thể thay đổi theo chỉ đạo của Lãnh đạo Bảo hiểm xã hội Việt Nam)",
        size=12, italic=True)

out_path = "/Users/hung/Projects/private/tieu_me/workspace/bao_cao/20260716_ke_hoach_hoi_nghi_truc_tuyen_TT12_chuanhoa_DM.docx"
doc.save(out_path)
print("SAVED:", out_path)
